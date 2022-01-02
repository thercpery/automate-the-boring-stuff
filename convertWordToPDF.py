import win32com.client
import docx
wordFilename = "word_example.docx"
pdfFilename = "pdf_example.pdf"

doc = docx.Document()
doc.add_heading("Header 0", 0)
doc.add_heading("Header 1", 1)
doc.add_heading("Header 2", 2)
doc.add_heading("Header 3", 3)
doc.add_heading("Header 4", 4)
doc.save(wordFilename)

wdFormatPDF = 17 # Word's numeric code for PDFs
wordObj = win32com.client.Dispatch("Word.Application")

docObj = wordObj.Documents.Open(wordFilename)
docObj.SaveAs(pdfFilename, FileFormat=wdFormatPDF)
docObj.Close()
wordObj.Quit()